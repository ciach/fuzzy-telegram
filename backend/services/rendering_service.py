from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML

from backend.schemas.contract_schema import ContractSchema
from backend.services.contract_service import ContractService


_CLAUSE_ORDINALS = {
    "clause_1": "PRIMERA",
    "clause_2": "SEGUNDA",
    "clause_3": "TERCERA",
    "clause_4": "CUARTA",
    "clause_5": "QUINTA",
}

_SELLER_PLACEHOLDER = "SELLER_PLACEHOLDER"
_BUYER_PLACEHOLDER = "BUYER_PLACEHOLDER"
_ADDRESS_PLACEHOLDER = "ADDRESS_PLACEHOLDER"
_PRICE_PLACEHOLDER = "PRICE_PLACEHOLDER"


class RenderingService:
    def __init__(self, contract_service: ContractService) -> None:
        self._contract_service = contract_service
        self._templates_dir = Path(__file__).resolve().parents[1] / "templates"
        self._jinja_env = Environment(
            loader=FileSystemLoader(self._templates_dir),
            autoescape=select_autoescape(["html", "xml"]),
            trim_blocks=True,
            lstrip_blocks=True,
        )

    def export_docx(self, contract_id: UUID) -> bytes:
        context = self._build_context(contract_id)
        document = Document()

        document.add_heading("CONTRATO DE COMPRAVENTA CON ARRAS PENITENCIALES", level=0)
        document.add_paragraph(f"En {context['metadata_location']}, a {context['metadata_date_text']}.")

        document.add_heading("REUNIDOS", level=1)
        document.add_paragraph(context["sellers_text"])
        document.add_paragraph(context["buyers_text"])
        document.add_paragraph(
            "Los comparecientes actúan en su propio nombre e interés y se reconocen capacidad legal suficiente."
        )

        document.add_heading("MANIFIESTAN", level=1)
        document.add_paragraph(f"Finca objeto de compraventa: {context['property_address']}")
        document.add_paragraph(f"Datos registrales: {context['property_registry']}")
        document.add_paragraph(f"Cargas: {context['property_cargas_text']}")
        document.add_paragraph(f"IBI anual aproximado: {context['property_ibi']} EUR")

        document.add_heading("CLÁUSULAS", level=1)
        for clause in context["clauses"]:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{clause['ordinal']}. ").bold = True
            paragraph.add_run(clause["text"])

        document.add_heading("CONDICIONES ECONÓMICAS", level=1)
        document.add_paragraph(f"Precio total: {context['total_price_text']} EUR")
        document.add_paragraph(f"Arras penitenciales: {context['arras_amount_text']} EUR")
        document.add_paragraph(f"Importe restante: {context['remaining_amount_text']} EUR")
        document.add_paragraph(f"Fecha límite de escritura: {context['deadline_text']}")

        document.add_heading("FIRMAS", level=1)
        document.add_paragraph("LAS COMPRADORAS")
        document.add_paragraph("LOS VENDEDORES")

        out = BytesIO()
        document.save(out)
        return out.getvalue()

    def export_pdf(self, contract_id: UUID) -> bytes:
        context = self._build_context(contract_id)
        template_name = _template_for_language(context["language"])
        template = self._jinja_env.get_template(template_name)
        html = template.render(**context)
        return HTML(string=html, base_url=str(self._templates_dir)).write_pdf()

    def _build_context(self, contract_id: UUID) -> dict[str, Any]:
        record = self._contract_service.get_contract(contract_id)
        schema = ContractSchema.model_validate(record.schema_json)

        sellers = schema.parties.sellers
        buyers = schema.parties.buyers

        sellers_text = _parties_text(sellers, fallback=_SELLER_PLACEHOLDER, name_placeholder=_SELLER_PLACEHOLDER)
        buyers_text = _parties_text(buyers, fallback=_BUYER_PLACEHOLDER, name_placeholder=_BUYER_PLACEHOLDER)

        sorted_clauses = [
            {
                "id": clause_id,
                "ordinal": _CLAUSE_ORDINALS.get(clause_id, clause_id.upper()),
                "text": text,
            }
            for clause_id, text in sorted(
                schema.clauses.items(),
                key=lambda item: item[0],
            )
            if text.strip()
        ]

        cargas_text = ", ".join(schema.property.cargas) if schema.property.cargas else "Sin cargas informadas"
        deadline_text = schema.financial.deadline.isoformat() if schema.financial.deadline else "Pendiente de definir"

        context = {
            "contract_id": str(contract_id),
            "language": schema.metadata.language,
            "metadata_location": schema.metadata.location,
            "metadata_date_text": schema.metadata.date.isoformat(),
            "sellers_text": sellers_text,
            "buyers_text": buyers_text,
            "property_address": schema.property.address.strip() or _ADDRESS_PLACEHOLDER,
            "property_registry": schema.property.registry.strip() or "Pendiente de completar",
            "property_cargas_text": cargas_text,
            "property_ibi": f"{schema.property.ibi:.2f}",
            "total_price_text": _price_text(schema.financial.total_price),
            "arras_amount_text": f"{schema.financial.arras_amount:.2f}",
            "remaining_amount_text": f"{schema.financial.remaining_amount:.2f}",
            "deadline_text": deadline_text,
            "clauses": sorted_clauses,
        }
        context["clauses"] = [
            {**item, "text": _apply_clause_placeholders(item["text"], context)} for item in sorted_clauses
        ]
        return context


def _parties_text(parties: list[Any], *, fallback: str, name_placeholder: str) -> str:
    if not parties:
        return fallback
    chunks: list[str] = []
    for party in parties:
        full_name = (party.full_name or "").strip()
        id_number = (party.id_number or "").strip()
        if not full_name and not id_number:
            continue
        id_text = f" (ID: {id_number})" if id_number else ""
        chunks.append(f"{full_name or name_placeholder}{id_text}")
    return "; ".join(chunks) if chunks else fallback


def _price_text(total_price: float) -> str:
    if total_price <= 0:
        return _PRICE_PLACEHOLDER
    return f"{total_price:.2f}"


def _template_for_language(language: str) -> str:
    if language == "ca":
        return "contract_ca.j2"
    if language == "en":
        return "contract_en.j2"
    return "contract_es.j2"


def _apply_clause_placeholders(clause_text: str, context: dict[str, Any]) -> str:
    replacements = {
        "{{buyers_names}}": context["buyers_text"],
        "{{sellers_names}}": context["sellers_text"],
        "{{total_price_eur}}": context["total_price_text"],
        "{{arras_amount_eur}}": context["arras_amount_text"],
        "{{remaining_amount_eur}}": context["remaining_amount_text"],
        "{{deadline_date}}": context["deadline_text"],
        "{{property_address}}": context["property_address"],
        "{{property_registry}}": context["property_registry"],
        "{{property_cargas}}": context["property_cargas_text"],
        "{{property_ibi_eur}}": context["property_ibi"],
        "{{location}}": context["metadata_location"],
        "{{contract_date}}": context["metadata_date_text"],
    }
    rendered = clause_text
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, str(value))
    return rendered
